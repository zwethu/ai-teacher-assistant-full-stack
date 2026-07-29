"""Export chat transcripts and individual responses as Markdown, PDF or DOCX.

Markdown is the canonical form: a chat is rendered to Markdown once and the PDF
and DOCX writers both consume that, so all three formats stay in step.

Both writers are pure Python (`markdown` + `xhtml2pdf` for PDF, `python-docx`
for DOCX). That is deliberate — the API image is `python:3.12-slim` with no apt
layer, and the usual HTML-to-PDF engines (WeasyPrint) need cairo/pango system
libraries that would have to be installed into it.
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any, Iterable

from docx import Document
from docx.shared import Pt, RGBColor

ExportFormat = str  # "markdown" | "pdf" | "docx"

SUPPORTED_FORMATS = ("markdown", "pdf", "docx")

MEDIA_TYPES = {
    "markdown": "text/markdown; charset=utf-8",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

EXTENSIONS = {"markdown": "md", "pdf": "pdf", "docx": "docx"}

# MILA tokens, inlined: the PDF is generated outside the browser so it cannot
# reach the design system's stylesheet.
_VIOLET = "#5f489c"
_INK = "#2a1e52"
_SLATE_700 = "#334155"
_SLATE_400 = "#94a3b8"
_SURFACE = "#f6f1fc"

_PDF_CSS = f"""
@page {{ size: A4; margin: 2cm 1.8cm; }}
body {{ font-family: Helvetica, Arial, sans-serif; font-size: 10.5pt;
        line-height: 1.5; color: {_SLATE_700}; }}
h1 {{ font-size: 18pt; color: {_INK}; margin: 0 0 4pt 0; }}
h2 {{ font-size: 13pt; color: {_INK}; margin: 14pt 0 4pt 0; }}
h3 {{ font-size: 11.5pt; color: {_INK}; margin: 12pt 0 3pt 0; }}
p {{ margin: 0 0 7pt 0; }}
ul, ol {{ margin: 0 0 7pt 16pt; }}
li {{ margin-bottom: 2pt; }}
code {{ font-family: Courier, monospace; font-size: 9.5pt; color: {_VIOLET};
        background-color: {_SURFACE}; }}
pre {{ font-family: Courier, monospace; font-size: 9pt; background-color: {_SURFACE};
       padding: 6pt; margin: 0 0 8pt 0; }}
blockquote {{ margin: 0 0 8pt 0; padding-left: 8pt; color: {_SLATE_400}; }}
table {{ width: 100%; margin: 0 0 8pt 0; }}
th {{ background-color: {_SURFACE}; color: {_INK}; text-align: left; padding: 4pt; }}
td {{ padding: 4pt; }}
.meta {{ color: {_SLATE_400}; font-size: 8.5pt; margin: 0 0 14pt 0; }}
.role {{ color: {_VIOLET}; font-size: 9pt; margin: 12pt 0 2pt 0; }}
"""


def _clean_filename(value: str, fallback: str = "chat") -> str:
    """Filesystem- and header-safe slug for Content-Disposition."""
    slug = re.sub(r"[^\w\s-]", "", value or "").strip()
    slug = re.sub(r"[\s_]+", "-", slug).strip("-")
    return (slug or fallback)[:60]


def _format_timestamp(value: Any) -> str:
    if not value:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d %b %Y, %H:%M")
    text = str(value)
    try:
        return datetime.fromisoformat(text.replace("Z", "+00:00")).strftime("%d %b %Y, %H:%M")
    except ValueError:
        return text


def _role_label(role: str) -> str:
    return "You" if role == "user" else "MILA"


def render_message_markdown(message: dict[str, Any]) -> str:
    """A single response, without the surrounding conversation."""
    return str(message.get("content") or "").strip()


def render_chat_markdown(chat: dict[str, Any], messages: Iterable[dict[str, Any]]) -> str:
    """The whole conversation, as Markdown."""
    title = str(chat.get("title") or "Chat")
    updated = _format_timestamp(chat.get("updated_at") or chat.get("created_at"))

    lines: list[str] = [f"# {title}", ""]
    if updated:
        lines += [f"_Last updated {updated} · exported from MILA_", ""]

    for message in messages:
        content = str(message.get("content") or "").strip()
        if not content:
            continue
        stamp = _format_timestamp(message.get("created_at"))
        heading = _role_label(str(message.get("role") or "user"))
        lines.append(f"## {heading}" + (f" · {stamp}" if stamp else ""))
        lines.append("")
        lines.append(content)
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def markdown_to_pdf(md_text: str, title: str) -> bytes:
    """Markdown → styled HTML → PDF."""
    import markdown as md_lib
    from xhtml2pdf import pisa

    body_html = md_lib.markdown(
        md_text,
        extensions=["extra", "sane_lists", "nl2br", "tables", "fenced_code"],
    )
    html = (
        "<html><head><meta charset='utf-8'>"
        f"<style>{_PDF_CSS}</style></head><body>{body_html}</body></html>"
    )

    buffer = io.BytesIO()
    result = pisa.CreatePDF(src=io.StringIO(html), dest=buffer, encoding="utf-8")
    if result.err:
        raise RuntimeError(f"PDF generation failed for {title!r}")
    return buffer.getvalue()


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_ORDERED_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*|_[^_]+_)")


def _add_runs(paragraph, text: str) -> None:
    """Split a line into runs so bold, italic and inline code survive."""
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0x5F, 0x48, 0x9C)
        elif len(part) > 2 and part[0] in "*_" and part[-1] == part[0]:
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(md_text: str, title: str) -> bytes:
    """Markdown → DOCX.

    A pragmatic subset — headings, bullet and numbered lists, fenced code and
    paragraphs, with bold/italic/code preserved inline. Enough for a readable
    Word document without pulling in a full Markdown AST.
    """
    document = Document()
    in_code = False
    code_lines: list[str] = []

    def flush_code() -> None:
        if not code_lines:
            return
        paragraph = document.add_paragraph()
        run = paragraph.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        code_lines.clear()

    for raw in md_text.splitlines():
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                flush_code()
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            continue

        heading = _HEADING_RE.match(line)
        if heading:
            document.add_heading(heading.group(2).strip(), level=min(len(heading.group(1)), 4))
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            _add_runs(document.add_paragraph(style="List Bullet"), bullet.group(1))
            continue

        ordered = _ORDERED_RE.match(line)
        if ordered:
            _add_runs(document.add_paragraph(style="List Number"), ordered.group(1))
            continue

        _add_runs(document.add_paragraph(), line)

    flush_code()

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_export(md_text: str, title: str, fmt: ExportFormat) -> tuple[bytes, str, str]:
    """Render `md_text` and return (payload, media_type, filename)."""
    if fmt not in SUPPORTED_FORMATS:
        raise ValueError(f"Unsupported export format: {fmt}")

    filename = f"{_clean_filename(title)}.{EXTENSIONS[fmt]}"

    if fmt == "markdown":
        return md_text.encode("utf-8"), MEDIA_TYPES[fmt], filename
    if fmt == "pdf":
        return markdown_to_pdf(md_text, title), MEDIA_TYPES[fmt], filename
    return markdown_to_docx(md_text, title), MEDIA_TYPES[fmt], filename
